#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
docx-table-title.py

テーブルのヘッダー行（最初の行）の背景色を深い青、文字色を白に変更するツール
"""

import sys
import os
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree


def set_cell_shading(cell, color):
    """セルの背景色を設定する"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 既存のshdを削除
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)

    # 新しいshdを作成
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_run_color(run, color):
    """runの文字色を設定する"""
    rPr = run._r.get_or_add_rPr()

    # 既存のcolorを削除
    for c in rPr.findall(qn('w:color')):
        rPr.remove(c)

    # 新しいcolorを作成
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), color)
    rPr.append(color_elem)


def set_paragraph_run_color(paragraph, color):
    """段落内のすべてのrunの文字色を設定する"""
    for run in paragraph.runs:
        set_run_color(run, color)

    # runがない場合でも、段落のrPrにデフォルト色を設定
    pPr = paragraph._p.get_or_add_pPr()

    # rPrを取得または作成
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)

    # 既存のcolorを削除
    for c in rPr.findall(qn('w:color')):
        rPr.remove(c)

    # 新しいcolorを作成
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), color)
    rPr.append(color_elem)


def process_table_headers(doc, header_bg_color, text_color):
    """すべてのテーブルのヘッダー行を処理する"""
    for table in doc.tables:
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # 背景色を設定
                set_cell_shading(cell, header_bg_color)

                # 各段落の文字色を設定
                for paragraph in cell.paragraphs:
                    set_paragraph_run_color(paragraph, text_color)


def main():
    if len(sys.argv) < 2:
        print("Usage: python docx-table-title.py <変換元ファイル名>")
        sys.exit(1)

    input_file = sys.argv[1]

    if not os.path.exists(input_file):
        print(f"Error: ファイルが見つかりません: {input_file}")
        sys.exit(1)

    # 出力先フォルダを作成
    input_dir = os.path.dirname(os.path.abspath(input_file))
    output_dir = os.path.join(input_dir, "changed")
    os.makedirs(output_dir, exist_ok=True)

    # 出力ファイルパス
    output_file = os.path.join(output_dir, os.path.basename(input_file))

    try:
        # ドキュメントを読み込む
        doc = Document(input_file)

        # テーブルヘッダーを処理
        # 深い青: 003366 または 1F4E79
        # 白: FFFFFF
        header_bg_color = "1F4E79"  # 深い青
        text_color = "FFFFFF"       # 白

        process_table_headers(doc, header_bg_color, text_color)

        # 保存
        doc.save(output_file)
        print(f"変換完了: {output_file}")

    except PermissionError as e:
        print(f"Error: ファイルへのアクセス権限がありません。ファイルが開かれている可能性があります。")
        print(f"詳細: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error: 処理中にエラーが発生しました: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
